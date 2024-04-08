
import docx

STOP_SYMBOL = ['/', ' ', '*', '+', '=', '-', ':', ';', '"', '%']
CALC_SYMBOL = ['/', '*', '+', '-', '**', '%']
dir_temp = 'temp.docx'


def read_doc():
    # получение списка уникальных параметров
    doc = docx.Document(dir_temp)
    len_paraf = len(doc.paragraphs)
    param_name_list = []
    for i in range(len_paraf):
        param_name = ''
        len_text = len(doc.paragraphs[i].text)
        for ii in range(len_text):
            read_symbol = doc.paragraphs[i].text[ii]
            if read_symbol in STOP_SYMBOL:
                if param_name != '':
                    param_name_list.append(param_name)
                    param_name = ''
            else:
                param_name += read_symbol
        if param_name != '':
            param_name_list.append(param_name)
    param_name_list_unique = sorted(set(param_name_list))
    print(param_name_list_unique)
    return param_name_list_unique

def write_doc(dict_param):
    doc_read = docx.Document(dir_temp)
    doc_write = docx.Document()
    doc_write.add_heading('Файл некого расчета!',)
    doc_write.add_paragraph('Формулы расчета:')
    for i in range(len(doc_read.paragraphs)):
        doc_write.add_paragraph(doc_read.paragraphs[i].text)
    doc_write.add_paragraph('')
    doc_write.add_paragraph('Введенные значения:')
    for key in dict_param:
        if dict_param[key] != None:
            doc_write.add_paragraph(f'{key}: {dict_param[key]}')
    doc_write.add_paragraph('')
    doc_write.add_paragraph('Полученные значения:')
    for i in range(len(doc_read.paragraphs)):
        param_name = ''
        str_paragraphs = ''
        key_calc = ''
        for ii in range(len(doc_read.paragraphs[i].text)):
            read_symbol = doc_read.paragraphs[i].text[ii]
            if read_symbol in STOP_SYMBOL:
                if read_symbol in CALC_SYMBOL:
                     str_paragraphs += read_symbol
                else:
                    if param_name != '':
                        if key_calc == '':
                            key_calc = param_name
                        else:
                            str_paragraphs += f' {dict_param[param_name]} '
                        param_name = ''
            else:
                param_name += read_symbol
        if param_name != '':
            str_paragraphs += f' {dict_param[param_name]} '
        try:
            calc_value = eval(str_paragraphs)
        except:
            print('ошибка')
        dict_param[key_calc] = calc_value
        # print(f'{key_calc} = {calc_value}')
        doc_write.add_paragraph(f'{key_calc}: {dict_param[key_calc]}')
    print(dict_param)
    doc_write.save('файл расчета.docx')
    return


if __name__ == '__main__':
    x = read_doc()
    print(x)
else:
    dir_temp = 'word/' + dir_temp